import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from AutoDocumentationCreator.const import *


def remove_suffix(title, suffix):
    if title.endswith(suffix):
        return title[:-len(suffix)]
    return title


class MsWordHandler:
    def __init__(self, ms_word_path):
        self.ms_word_path = ms_word_path
        self.doc = docx.Document(self.ms_word_path)

    @classmethod
    def remove_suffix(self, s, suffix):
        if s.endswith(suffix):
            return s[:-len(suffix)]
        return s

    @classmethod
    def extract_slice_name(self, image_path):
        try:
            # net = image_path.name.split('_')
            net = image_path.split('_')
            if len(net) == INDEX_OF_ITEM_LIST:
                title = net[INDEX_OF_LAST_ITEM]
            else:
                title = (net[INDEX_OF_SECOND_LAST_ITEM] + ':' + net[INDEX_OF_LAST_ITEM])
            title_final = remove_suffix(title, '].png')

            return title_final
        except Exception as e:
            print("Exception occurred:", e)

    def attach_images_to_ms_word(self, image_path, index):
        try:
            self.doc.add_picture(image_path, width=Cm(IMG_WIDTH_CM))
            first_paragraph = self.doc.paragraphs[-1]
            first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            extract_name = self.extract_slice_name(image_path)

            paragra = self.doc.add_paragraph()
            paragra.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragra.add_run('Figure' + " " + str(index) + "." + extract_name).bold = True
            print(index)
            self.doc.save(self.ms_word_path)
        except Exception as e:
            print("Exception occurred:", e)
            self.doc.save(self.ms_word_path)


